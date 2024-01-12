from docx import Document
from docx2pdf import convert
import os
import glob

def replace_text_in_docx(docx_path, old_text, new_text):
    doc = Document(docx_path)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace(old_text, new_text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace(old_text, new_text)

    doc.save("temp.docx")

def convert_docx_to_pdf(docx_path, pdf_path):
    convert(docx_path, pdf_path)

if __name__ == "__main__":
    input_docx = glob.glob('*.docx') # will find any docx files
    print(input_docx)
    position_text = "[Position]"
    new_position = input("Position title: ")
    company_text = "[Company]"
    new_company = input("Company name: ")
    output_pdf = f"CoverLetter_{new_company}.pdf"  # output cover letter as pdf

    for elem in input_docx:
        replace_text_in_docx(elem, position_text, new_position)
        replace_text_in_docx("temp.docx", company_text, new_company)
    convert_docx_to_pdf("temp.docx", output_pdf)
    os.remove("temp.docx")
    print(f"Conversion complete. PDF saved at: {output_pdf}")
