import os
import string
import glob
import base64
import requests
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email import encoders
from docx import Document
from docx2pdf import convert
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from googleapiclient.discovery import build
from requests import HTTPError
from oauth2client import client
from oauth2client import tools
from oauth2client.file import Storage

def replace_text_in_docx(docx_path, old_text, new_text):
    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace(old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace(old_text, new_text)
    doc.save("temp.docx")

def convert_docx_to_pdf(docx_path, pdf_path):
    convert(docx_path, pdf_path)

def get_job_details(url):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        company_title = soup.find('span', class_='business').find(['strong', 'a']).text.strip()
        position_title = soup.select_one('h1.title span[property="title"]').text.strip()
        return company_title.translate(str.maketrans('', '', string.punctuation)), position_title.title()
    else:
        return "Failed to fetch the URL. Please check if it's valid."

def get_email_after_click(url):
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--log-level=3')
    driver = webdriver.Chrome(options=options)
    driver.get(url)
    try:
        apply_button = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, "applynowbutton")))
        apply_button.click()
        email_element = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, "#howtoapply a[href^='mailto:']")))
        email_address = email_element.get_attribute("href").split(":")[1]
        return email_address
    finally:
        driver.quit()

def send_email(recipient_email, subject, email_content, output_pdf):
    SCOPES = ["https://www.googleapis.com/auth/gmail.send"]
    creds_path = os.path.join('./', 'credential_sample.json')
    store = Storage(creds_path)
    credentials = store.get()
    if not credentials or credentials.invalid:
        flow = client.flow_from_clientsecrets('credentials.json', SCOPES)
        credentials = tools.run_flow(flow, store)
    service = build('gmail', 'v1', credentials=credentials)

    with open("Leitch, Mason.pdf", "rb") as attachment:
        resume = MIMEBase("application", "octet-stream")
        resume.set_payload(attachment.read())
    encoders.encode_base64(resume)
    resume.add_header("Content-Disposition", f"attachment; filename=Leitch, Mason.pdf")

    filename = os.path.basename(output_pdf)
    with open(output_pdf, "rb") as attachment:
        cover_letter = MIMEBase("application", "octet-stream")
        cover_letter.set_payload(attachment.read())
    encoders.encode_base64(cover_letter)
    cover_letter.add_header("Content-Disposition", f"attachment; filename={filename}")

    message = MIMEMultipart()
    text = MIMEText(email_content, 'html')
    message['To'] = recipient_email
    message['Subject'] = subject
    message.attach(text)
    message.attach(resume)
    message.attach(cover_letter)
    create_message = {'raw': base64.urlsafe_b64encode(message.as_bytes()).decode('utf-8','ignore')}

    try:
        message = (service.users().messages().send(userId="me", body=create_message).execute())
        print(f'Sent message to {recipient_email} Message Id: {message["id"]}')
    except HTTPError as error:
        print(f'An error occurred: {error}')
        message = None

def parse_and_replace():
    url = input("Enter the URL: ")
    if url == 'exit':
        exit()
    company_title, position_title = get_job_details(url)
    email_address = get_email_after_click(url)
    if email_address:
        print("Email address found:", email_address)
    else:
        print("Email address not found.")

    input_docx = glob.glob('*.docx')
    print("Modifying ", input_docx)
    
    position_text = "[Position]"
    new_position = position_title

    company_text = "[Company]"
    new_company = company_title
    if os.path.exists("temp.docx"):
        os.remove("temp.docx")
    if not os.path.exists(f"./Company Cover Letters/{new_company}"):
        os.makedirs(f"./Company Cover Letters/{new_company}")
    output_pdf = f"./Company Cover Letters/{new_company}/CoverLetter_{new_company}.pdf"
    
    for elem in input_docx:
        replace_text_in_docx(elem, position_text, new_position)
        replace_text_in_docx("temp.docx", company_text, new_company)
    convert_docx_to_pdf("temp.docx", output_pdf)
    os.remove("temp.docx")

    with open('email.txt', 'r') as file:
        content = file.read()
    content = content.replace('[Position]', new_position)
    content = content.replace('[Company]', new_company)
    subject = f"Application for {new_position} at {new_company}"

    print(f"Conversion complete. PDF saved at: {output_pdf}")

    return content, output_pdf, subject, email_address

def main():
    while True:
        content, output_pdf, subject, recipient_email = parse_and_replace()
        send_email(recipient_email, subject, content, output_pdf)
        print("Email sent.")

if __name__ == "__main__":
    main()
